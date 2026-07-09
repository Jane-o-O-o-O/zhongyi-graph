from __future__ import annotations

import csv
import json
import re
from dataclasses import dataclass
from io import BytesIO, StringIO
from pathlib import Path

from app.models.ingestion import DocumentChunk, DocumentPage, ExtractionUnit, SourceManifest

TEXT_MIME_TYPES = {"text/plain", "text/markdown", "application/x-markdown"}
CSV_MIME_TYPES = {"text/csv", "application/csv"}
JSON_MIME_TYPES = {"application/json"}
IMAGE_MIME_PREFIX = "image/"

STRUCTURAL_LINE_RE = re.compile(r"^<(目录|篇名)>(.*)$")
BODY_PREFIX_RE = re.compile(r"^(属性|内容)：")
SENTENCE_BOUNDARY_RE = re.compile(r"(?<=[。！？；;])")


@dataclass(frozen=True)
class ParsedText:
    text: str
    content_type: str
    status: str


@dataclass(frozen=True)
class SectionBlock:
    title: str
    directory: str
    content: str
    char_start: int
    char_end: int
    unit_type: str


class DocumentParser:
    def __init__(
        self,
        ocr_client=None,
        chunk_chars: int | None = None,
        parent_target_chars: int = 2400,
        parent_max_chars: int = 4000,
        child_target_chars: int = 360,
        child_max_chars: int = 520,
    ):
        self.ocr_client = ocr_client
        self.parent_target_chars = parent_target_chars
        self.parent_max_chars = parent_max_chars
        self.child_target_chars = chunk_chars or child_target_chars
        self.child_max_chars = max(child_max_chars, self.child_target_chars)

    def parse(
        self,
        source: SourceManifest,
        content: bytes,
    ) -> tuple[list[DocumentPage], list[ExtractionUnit], list[DocumentChunk], str]:
        parsed = self._parse_text(source, content)
        if not parsed.text.strip():
            return [], [], [], parsed.status

        text = _normalize_text(parsed.text)
        page = DocumentPage(
            page_id=f"page:{source.source_id}:1",
            source_id=source.source_id,
            page_number=1,
            text=text,
        )
        units = self._build_extraction_units(
            source=source,
            page_id=page.page_id,
            text=text,
            content_type=parsed.content_type,
        )
        chunks = self._build_child_chunks(
            source_id=source.source_id,
            units=units,
            content_type=parsed.content_type,
        )
        return [page], units, chunks, parsed.status

    def _parse_text(self, source: SourceManifest, content: bytes) -> ParsedText:
        content_type = "text"
        status = "parsed"
        if source.mime_type in TEXT_MIME_TYPES or _suffix(source.filename) in {".txt", ".md"}:
            text = _decode(content)
        elif source.mime_type in CSV_MIME_TYPES or _suffix(source.filename) == ".csv":
            text = _csv_to_markdown(content)
            content_type = "table"
        elif source.mime_type in JSON_MIME_TYPES or _suffix(source.filename) == ".json":
            text = json.dumps(json.loads(_decode(content)), ensure_ascii=False, indent=2)
            content_type = "json"
        elif source.mime_type.startswith(IMAGE_MIME_PREFIX):
            if not self.ocr_client:
                return ParsedText("", content_type, "requires_ocr")
            text = self.ocr_client.recognize_image(content, mime_type=source.mime_type)
            content_type = "ocr"
        elif _suffix(source.filename) == ".docx":
            text = _docx_text(content)
        elif _suffix(source.filename) == ".pdf":
            text = _pdf_text(content)
            if not text.strip():
                status = "requires_ocr"
        else:
            text = _decode(content)
        return ParsedText(text=text, content_type=content_type, status=status)

    def _build_extraction_units(
        self,
        *,
        source: SourceManifest,
        page_id: str,
        text: str,
        content_type: str,
    ) -> list[ExtractionUnit]:
        blocks = _extract_section_blocks(
            text,
            fallback_title=Path(source.filename).stem,
            fallback_unit_type=_unit_type_for_content(content_type),
        )
        if not blocks:
            blocks = [
                SectionBlock(
                    title=Path(source.filename).stem,
                    directory="",
                    content=text.strip(),
                    char_start=0,
                    char_end=len(text),
                    unit_type=_unit_type_for_content(content_type),
                )
            ]

        units: list[ExtractionUnit] = []
        for block in blocks:
            section_path = _section_path(Path(source.filename).stem, block.directory, block.title)
            for content, start, end in _semantic_splits(
                block.content,
                target_chars=self.parent_target_chars,
                max_chars=self.parent_max_chars,
                absolute_start=block.char_start,
            ):
                if not content:
                    continue
                index = len(units) + 1
                title = block.title or Path(source.filename).stem
                if len(block.content.strip()) > self.parent_max_chars:
                    title = f"{title} #{index}"
                units.append(
                    ExtractionUnit(
                        unit_id=f"unit:{source.source_id}:{index:04d}",
                        source_id=source.source_id,
                        page_id=page_id,
                        unit_index=index,
                        title=title,
                        content=content,
                        unit_type=block.unit_type,
                        section_path=section_path,
                        token_count=len(content),
                        char_start=start,
                        char_end=end,
                        metadata={
                            "directory": block.directory,
                            "source_filename": source.filename,
                            "split_strategy": "structural_sentence",
                        },
                    )
                )
        return units

    def _build_child_chunks(
        self,
        *,
        source_id: str,
        units: list[ExtractionUnit],
        content_type: str,
    ) -> list[DocumentChunk]:
        chunks: list[DocumentChunk] = []
        for unit in units:
            for content, start, end in _semantic_splits(
                unit.content,
                target_chars=self.child_target_chars,
                max_chars=self.child_max_chars,
                absolute_start=unit.char_start,
            ):
                if not content:
                    continue
                chunk_index = len(chunks) + 1
                chunks.append(
                    DocumentChunk(
                        chunk_id=f"chunk:{source_id}:{chunk_index:04d}",
                        source_id=source_id,
                        page_id=unit.page_id,
                        chunk_index=chunk_index,
                        content=content,
                        parent_unit_id=unit.unit_id,
                        unit_index=unit.unit_index,
                        content_type=content_type,
                        section_title=unit.title,
                        token_count=len(content),
                        char_start=start,
                        char_end=end,
                        metadata={
                            "parent_unit_id": unit.unit_id,
                            "parent_unit_index": unit.unit_index,
                            "parent_unit_title": unit.title,
                            "section_path": unit.section_path,
                            "split_strategy": "sentence_window",
                        },
                    )
                )
        return chunks


def _extract_section_blocks(
    text: str,
    fallback_title: str,
    fallback_unit_type: str,
) -> list[SectionBlock]:
    lines = text.splitlines(keepends=True)
    current_directory = ""
    current_title = ""
    document_title = fallback_title
    active: dict | None = None
    blocks: list[SectionBlock] = []
    offset = 0

    for line in lines:
        stripped = line.strip()
        structural_match = STRUCTURAL_LINE_RE.match(stripped)
        if structural_match:
            marker, value = structural_match.groups()
            value = value.strip()
            if marker == "目录":
                current_directory = value
            elif marker == "篇名":
                if not current_title and not blocks and active is None:
                    document_title = value or document_title
                current_title = value or document_title
            offset += len(line)
            continue

        if BODY_PREFIX_RE.match(stripped):
            if active:
                _finish_block(active, offset, blocks)
            active = {
                "title": current_title or document_title,
                "directory": current_directory,
                "parts": [_strip_body_prefix(line)],
                "start": offset,
                "unit_type": "article",
            }
            offset += len(line)
            continue

        if active:
            active["parts"].append(line)
        offset += len(line)

    if active:
        _finish_block(active, len(text), blocks)

    if blocks:
        return [block for block in blocks if block.content.strip()]

    stripped = text.strip()
    if not stripped:
        return []
    return [
        SectionBlock(
            title=document_title,
            directory=current_directory,
            content=stripped,
            char_start=text.find(stripped),
            char_end=text.find(stripped) + len(stripped),
            unit_type=fallback_unit_type,
        )
    ]


def _finish_block(active: dict, end: int, blocks: list[SectionBlock]) -> None:
    content = "".join(active["parts"]).strip()
    if not content:
        return
    blocks.append(
        SectionBlock(
            title=active["title"],
            directory=active["directory"],
            content=content,
            char_start=active["start"],
            char_end=end,
            unit_type=active["unit_type"],
        )
    )


def _strip_body_prefix(line: str) -> str:
    return BODY_PREFIX_RE.sub("", line, count=1).strip()


def _semantic_splits(
    text: str,
    *,
    target_chars: int,
    max_chars: int,
    absolute_start: int,
) -> list[tuple[str, int, int]]:
    normalized = text.strip()
    if not normalized:
        return []
    if len(normalized) <= target_chars:
        start = text.find(normalized)
        if start < 0:
            start = 0
        return [(normalized, absolute_start + start, absolute_start + start + len(normalized))]

    pieces = _sentence_pieces(normalized)
    splits: list[tuple[str, int, int]] = []
    buffer = ""
    cursor = 0
    split_start = 0

    for piece in pieces:
        candidate = _join_piece(buffer, piece)
        if buffer and len(candidate) >= target_chars:
            split_text = buffer.strip()
            start = normalized.find(split_text, split_start)
            if start < 0:
                start = cursor
            end = start + len(split_text)
            splits.append((split_text, absolute_start + start, absolute_start + end))
            split_start = end
            cursor = end
            buffer = piece.strip()
            continue
        if len(candidate) > max_chars:
            for hard_piece in _hard_wrap(candidate, max_chars):
                split_text = hard_piece.strip()
                if not split_text:
                    continue
                start = normalized.find(split_text, split_start)
                if start < 0:
                    start = cursor
                end = start + len(split_text)
                splits.append((split_text, absolute_start + start, absolute_start + end))
                split_start = end
                cursor = end
            buffer = ""
            continue
        buffer = candidate

    if buffer.strip():
        split_text = buffer.strip()
        start = normalized.find(split_text, split_start)
        if start < 0:
            start = cursor
        splits.append((split_text, absolute_start + start, absolute_start + start + len(split_text)))

    return splits


def _sentence_pieces(text: str) -> list[str]:
    pieces: list[str] = []
    for paragraph in re.split(r"\n{2,}", text):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        sentence_parts = [part.strip() for part in SENTENCE_BOUNDARY_RE.split(paragraph) if part.strip()]
        pieces.extend(sentence_parts or [paragraph])
    return pieces or [text]


def _join_piece(left: str, right: str) -> str:
    if not left:
        return right.strip()
    if left.endswith("\n"):
        return f"{left}{right.strip()}"
    return f"{left}{right.strip()}"


def _hard_wrap(text: str, max_chars: int) -> list[str]:
    return [text[start : start + max_chars] for start in range(0, len(text), max_chars)]


def _merge_tiny_tail(
    splits: list[tuple[str, int, int]],
    *,
    max_chars: int,
) -> list[tuple[str, int, int]]:
    if len(splits) < 2:
        return splits
    previous_text, previous_start, _previous_end = splits[-2]
    tail_text, _tail_start, tail_end = splits[-1]
    if len(tail_text) >= 40 or len(previous_text) + len(tail_text) > max_chars:
        return splits
    merged = (previous_text + tail_text, previous_start, tail_end)
    return splits[:-2] + [merged]


def _section_path(source_title: str, directory: str, title: str) -> list[str]:
    parts = [source_title]
    if directory:
        parts.append(directory)
    if title and title not in parts:
        parts.append(title)
    return parts


def _unit_type_for_content(content_type: str) -> str:
    return {
        "table": "table",
        "json": "json",
        "ocr": "ocr",
    }.get(content_type, "text")


def _normalize_text(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()


def _decode(content: bytes) -> str:
    for encoding in ("utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _csv_to_markdown(content: bytes) -> str:
    text = _decode(content)
    rows = list(csv.reader(StringIO(text)))
    if not rows:
        return ""
    return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)


def _docx_text(content: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _pdf_text(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _suffix(filename: str) -> str:
    return Path(filename).suffix.lower()
